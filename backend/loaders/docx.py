"""
DOCX Loader Module

Extracts text from Microsoft Word documents.
"""

from docx import Document as DocxDocument
from langchain_core.documents import Document
from typing import List


def load_docx(file_path: str) -> List[Document]:
    """
    Load a DOCX file and return a single Document.
    
    Unlike PDFs, we treat the entire DOCX as one document
    since Word docs don't have strict page boundaries.
    """
    documents = []
    
    try:
        doc = DocxDocument(file_path)
        print(f"📝 Loading DOCX: {file_path}")
        
        # Extract all paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        content = "\n".join(full_text)
        
        if content:
            documents.append(Document(
                page_content=content,
                metadata={
                    "source": file_path,
                    "paragraphs": len(full_text)
                }
            ))
            print(f"   ✅ Extracted {len(full_text)} paragraphs")
        
    except Exception as e:
        print(f"   ❌ Error loading DOCX {file_path}: {e}")
        
    return documents
